from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.db.models import Avg
from students.models import Student, Enrollment
from complaints.models import Complaint
from semesters.models import Semester
from accounts.models import User
import io, os, openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn; from docx.oxml import OxmlElement

BASE_DIR  = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
LOGO_PATH = os.path.join(BASE_DIR,'static','images','logo.jpg')
NAVY=RGBColor(0x1B,0x2A,0x5E); GOLD=RGBColor(0xC9,0xA8,0x4C)

def _sem():
    return Semester.objects.filter(is_active=True).order_by('-created_at').first()

def _rtl(p):
    pPr=p._p.get_or_add_pPr(); b=OxmlElement('w:bidi'); b.set(qn('w:val'),'1'); pPr.insert(0,b)

def _shd(cell,fill):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),fill); tcPr.append(s)

def _no_border(cell):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcB=OxmlElement('w:tcBorders')
    for side in['top','left','bottom','right']:
        b=OxmlElement(f'w:{side}'); b.set(qn('w:val'),'none'); tcB.append(b)
    tcPr.append(tcB)

def _gold_border(p,position='bottom'):
    pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
    b=OxmlElement(f'w:{position}'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'12')
    b.set(qn('w:space'),'1'); b.set(qn('w:color'),'C9A84C'); pBdr.append(b); pPr.append(pBdr)

def add_hdr_ftr(doc, advisor_name):
    for section in doc.sections:
        section.header_distance=Cm(0.5); section.footer_distance=Cm(0.5)
        hdr=section.header; hdr.is_linked_to_previous=False
        for p in hdr.paragraphs: p.clear()
        ht=hdr.add_table(1,3,width=Cm(17)); ht.style='Table Grid'; ht.autofit=False
        for row in ht.rows:
            for cell in row.cells: _no_border(cell)
        lc=ht.cell(0,0); lc.width=Cm(3); lp=lc.paragraphs[0]; lp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=lp.add_run()
        if os.path.exists(LOGO_PATH): run.add_picture(LOGO_PATH,width=Cm(2.5))
        nc=ht.cell(0,1); nc.width=Cm(11); nc.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
        for i,(txt,sz,color,bold) in enumerate([('جامعة المنصورة الأهلية',14,NAVY,True),('كلية الهندسة',12,GOLD,True),('نظام الإرشاد الأكاديمي',10,NAVY,False)]):
            np=nc.paragraphs[0] if i==0 else nc.add_paragraph()
            np.alignment=WD_ALIGN_PARAGRAPH.CENTER; _rtl(np)
            r=np.add_run(txt); r.font.bold=bold; r.font.size=Pt(sz); r.font.color.rgb=color; r.font.name='Times New Roman'
        ht.cell(0,2).width=Cm(3)
        _gold_border(hdr.add_paragraph())
        ftr=section.footer; ftr.is_linked_to_previous=False
        for p in ftr.paragraphs: p.clear()
        fdiv=ftr.paragraphs[0]; _gold_border(fdiv,'top')
        ft=ftr.add_table(1,2,width=Cm(17)); ft.style='Table Grid'; ft.autofit=False
        for row in ft.rows:
            for cell in row.cells: _no_border(cell)
        for cell,title,sub in[(ft.cell(0,0),'توقيع المرشد الأكاديمي',advisor_name),(ft.cell(0,1),'توقيع الطالب','الاسم: ___________________')]:
            cell.width=Cm(8.5)
            for j,(txt,bold,sz) in enumerate([(title,True,10),(sub,False,9),('التوقيع: ___________________',False,10)]):
                pp=cell.paragraphs[0] if j==0 else cell.add_paragraph()
                pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; _rtl(pp)
                rr=pp.add_run(txt); rr.font.bold=bold; rr.font.size=Pt(sz); rr.font.color.rgb=NAVY; rr.font.name='Times New Roman'

def _sec(doc,text):
    p=doc.add_paragraph(); _rtl(p); r=p.add_run(f'  {text}  ')
    r.font.bold=True; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); r.font.name='Times New Roman'
    pPr=p._p.get_or_add_pPr(); s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),'1B2A5E'); pPr.append(s)
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(4)

def _row(doc,label,value):
    t=doc.add_table(1,2); t.autofit=False; t.style='Table Grid'
    c1=t.cell(0,0); c1.width=Cm(4); p1=c1.paragraphs[0]; _rtl(p1); p1.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r1=p1.add_run(label); r1.font.bold=True; r1.font.size=Pt(10); r1.font.name='Times New Roman'; _shd(c1,'E8ECF5')
    c2=t.cell(0,1); c2.width=Cm(12.5); p2=c2.paragraphs[0]; _rtl(p2); p2.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r2=p2.add_run(str(value) if value else '-'); r2.font.size=Pt(10); r2.font.name='Times New Roman'
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

@login_required
def reports_home(request):
    return render(request,'reports/home.html',{'semester':_sem()})

@login_required
def student_report_word(request, pk):
    student=get_object_or_404(Student,pk=pk); sem=_sem()
    enr=Enrollment.objects.filter(student=student,semester=sem).first()
    doc=Document()
    for s in doc.sections:
        s.page_width=Cm(21); s.page_height=Cm(29.7)
        s.top_margin=Cm(3.5); s.bottom_margin=Cm(3); s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)
    add_hdr_ftr(doc, str(student.advisor or request.user))
    t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER; _rtl(t)
    tr=t.add_run('استمارة الإرشاد الأكاديمي'); tr.font.bold=True; tr.font.size=Pt(18); tr.font.color.rgb=NAVY; tr.font.name='Times New Roman'
    sb=doc.add_paragraph(); sb.alignment=WD_ALIGN_PARAGRAPH.CENTER; _rtl(sb)
    sr=sb.add_run(f'{sem} - {sem.academic_year}' if sem else ''); sr.font.size=Pt(12); sr.font.color.rgb=GOLD; sr.font.name='Times New Roman'; sb.paragraph_format.space_after=Pt(12)
    _sec(doc,'بيانات الطالب')
    _row(doc,'الاسم الكامل:',student.full_name); _row(doc,'الرقم الجامعي:',student.student_id)
    _row(doc,'الرقم القومي:',student.national_id); _row(doc,'السن:',f'{student.age} سنة' if student.age else '-')
    _row(doc,'المستوى:',student.get_level_display()); _row(doc,'الساعات المنتهية:',f'{student.completed_hours} ساعة')
    if enr:
        _row(doc,'الساعات المسجلة:',f'{enr.registered_hours} ساعة')
        _row(doc,'المعدل التراكمي:',str(enr.gpa) if enr.gpa else '-')
        if enr.courses.exists():
            _sec(doc,'المواد المسجلة')
            courses=list(enr.courses.all())
            ct=doc.add_table(1+len(courses),3); ct.style='Table Grid'; ct.autofit=False
            for i,h in enumerate(['اسم المادة','كود المادة','الساعات']):
                cell=ct.cell(0,i); p=cell.paragraphs[0]; _rtl(p); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                r=p.add_run(h); r.font.bold=True; r.font.size=Pt(10); r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _shd(cell,'1B2A5E')
            for ri,c in enumerate(courses,1):
                for ci,val in enumerate([c.course_name,c.course_code or '-',str(c.credit_hours)]):
                    cell=ct.cell(ri,ci); p=cell.paragraphs[0]; _rtl(p); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    r=p.add_run(val); r.font.size=Pt(10); r.font.name='Times New Roman'
        if enr.semester_goal: _sec(doc,'هدف الفصل'); doc.add_paragraph(enr.semester_goal)
        if enr.advisor_notes: _sec(doc,'ملاحظات المرشد'); doc.add_paragraph(enr.advisor_notes)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    resp=HttpResponse(buf.read(),content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    resp['Content-Disposition']=f'attachment; filename="student_{student.student_id}.docx"'; return resp

@login_required
def semester_excel(request):
    sem=_sem()
    if not sem: return HttpResponse('لا يوجد فصل دراسي نشط')
    enrs=Enrollment.objects.filter(semester=sem).select_related('student','student__advisor')
    comps=Complaint.objects.filter(semester=sem).select_related('student')
    wb=openpyxl.Workbook()
    NAV='FF1B2A5E'; GLD='FFC9A84C'; WHT='FFFFFFFF'; LBL='FFE8ECF5'
    hf=Font(name='Cairo',bold=True,color=WHT,size=11)
    hfill=PatternFill('solid',fgColor=NAV)
    hal=Alignment(horizontal='center',vertical='center',wrap_text=True)
    thin=Side(style='thin',color='DDDDDD'); bdr=Border(left=thin,right=thin,top=thin,bottom=thin)
    def hrow(ws,row,cols):
        for c in range(1,cols+1): ws.cell(row,c).font=hf; ws.cell(row,c).fill=hfill; ws.cell(row,c).alignment=hal; ws.cell(row,c).border=bdr
    def drow(ws,row,cols,alt=False):
        fill=PatternFill('solid',fgColor='FFF8F9FC' if alt else WHT)
        for c in range(1,cols+1): ws.cell(row,c).font=Font(name='Cairo',size=10); ws.cell(row,c).fill=fill; ws.cell(row,c).alignment=Alignment(horizontal='center',vertical='center'); ws.cell(row,c).border=bdr
    # Sheet 1: Summary
    ws1=wb.active; ws1.title='ملخص الفصل'; ws1.sheet_view.rightToLeft=True
    ws1.column_dimensions['A'].width=38; ws1.column_dimensions['B'].width=18
    avg=enrs.aggregate(a=Avg('gpa'))['a'] or 0
    rows=[('البيان','القيمة'),('الفصل الدراسي',str(sem)),('إجمالي الطلاب',enrs.count()),('متوسط المعدل التراكمي',round(float(avg),2)),('الطلاب المتفوقون (GPA≥3)',enrs.filter(gpa__gte=3).count()),('الطلاب المستقرون (2.5≤GPA<3)',enrs.filter(gpa__gte=2.5,gpa__lt=3).count()),('الطلاب في خطر (GPA<2.5)',enrs.filter(gpa__lt=2.5).count()),('إجمالي الشكاوى',comps.count()),('شكاوى مغلقة',comps.filter(status='closed').count()),('شكاوى مصعّدة',comps.filter(status='escalated').count()),('معدل الإقرار بالبيانات',f"{round(enrs.filter(data_confirmed=True).count()/max(enrs.count(),1)*100,1)}%")]
    for ri,(a,b) in enumerate(rows,1):
        ws1.cell(ri,1,a); ws1.cell(ri,2,b); ws1.row_dimensions[ri].height=22
        if ri==1: hrow(ws1,1,2)
        else: ws1.cell(ri,1).font=Font(name='Cairo',bold=True,size=10); ws1.cell(ri,1).fill=PatternFill('solid',fgColor=LBL); ws1.cell(ri,1).alignment=Alignment(horizontal='right',vertical='center'); ws1.cell(ri,2).font=Font(name='Cairo',size=10); ws1.cell(ri,2).alignment=Alignment(horizontal='center',vertical='center')
        for c in[1,2]: ws1.cell(ri,c).border=bdr
    # Sheet 2: Students
    ws2=wb.create_sheet('بيانات الطلاب'); ws2.sheet_view.rightToLeft=True
    h2=['الرقم الجامعي','الاسم الكامل','المرشد الأكاديمي','المستوى','الساعات المنتهية','الساعات المسجلة','المعدل','الحالة','EWS','تأكيد البيانات','هدف الفصل']
    w2=[16,28,26,16,16,16,12,14,10,16,35]
    for i,(h,w) in enumerate(zip(h2,w2),1): ws2.cell(1,i,h); ws2.column_dimensions[get_column_letter(i)].width=w
    hrow(ws2,1,len(h2)); ws2.row_dimensions[1].height=26
    for ri,e in enumerate(enrs,2):
        g=float(e.gpa) if e.gpa else None
        status='متفوق' if g and g>=3 else 'مستقر' if g and g>=2.5 else 'في خطر' if g else '-'
        vals=[e.student.student_id,e.student.full_name,e.student.advisor.display_name if e.student.advisor else '-',e.student.get_level_display(),e.student.completed_hours,e.registered_hours,g or '-',status,e.early_warning_score,'✓' if e.data_confirmed else '✗',e.semester_goal[:50] if e.semester_goal else '']
        for ci,val in enumerate(vals,1): ws2.cell(ri,ci,val)
        drow(ws2,ri,len(h2),ri%2==0); ws2.row_dimensions[ri].height=20
        sc=ws2.cell(ri,8)
        if status=='متفوق': sc.font=Font(name='Cairo',bold=True,color='FF10b981',size=10)
        elif status=='مستقر': sc.font=Font(name='Cairo',bold=True,color='FF3b82f6',size=10)
        elif status=='في خطر': sc.font=Font(name='Cairo',bold=True,color='FFef4444',size=10)
    # Sheet 3: Courses
    ws3=wb.create_sheet('المواد المسجلة'); ws3.sheet_view.rightToLeft=True
    h3=['الرقم الجامعي','اسم الطالب','اسم المادة','كود المادة','الساعات']
    for i,(h,w) in enumerate(zip(h3,[16,28,35,16,10]),1): ws3.cell(1,i,h); ws3.column_dimensions[get_column_letter(i)].width=w
    hrow(ws3,1,5); ri=2
    for e in enrs:
        for c in e.courses.all(): ws3.cell(ri,1,e.student.student_id); ws3.cell(ri,2,e.student.full_name); ws3.cell(ri,3,c.course_name); ws3.cell(ri,4,c.course_code or '-'); ws3.cell(ri,5,c.credit_hours); drow(ws3,ri,5,ri%2==0); ws3.row_dimensions[ri].height=20; ri+=1
    # Sheet 4: Complaints
    ws4=wb.create_sheet('الشكاوى'); ws4.sheet_view.rightToLeft=True
    h4=['الرقم الجامعي','الطالب','عنوان الشكوى','وصف المشكلة','الحالة','مُصعَّدة؟','تاريخ التسجيل']
    for i,(h,w) in enumerate(zip(h4,[16,26,28,40,14,12,16]),1): ws4.cell(1,i,h); ws4.column_dimensions[get_column_letter(i)].width=w
    hrow(ws4,1,7)
    for ri,c in enumerate(comps,2):
        vals=[c.student.student_id,c.student.full_name,c.title,c.student_description,c.get_status_display(),'نعم' if c.status=='escalated' else 'لا',c.created_at.strftime('%Y-%m-%d')]
        for ci,val in enumerate(vals,1): ws4.cell(ri,ci,val)
        drow(ws4,ri,7,ri%2==0); ws4.row_dimensions[ri].height=20
    buf=io.BytesIO(); wb.save(buf); buf.seek(0)
    resp=HttpResponse(buf.read(),content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    resp['Content-Disposition']=f'attachment; filename="mnu_{sem.academic_year}_{sem.semester_type}.xlsx"'; return resp

